# =============================================================================
# anonimizar.py — Anonimizador multi-jurisdiccional v2.2
#
# Jurisdicciones soportadas:
#   rgpd      — UE (RGPD/GDPR 2016/679)
#   chile     — Ley 21.719 / Ley 19.628 (Chile)
#   brasil    — LGPD Lei 13.709/2018 (Brasil)
#   mexico    — LFPDPPP 2010/2025 (México)
#   colombia  — Ley 1581/2012 + Decreto 1377/2013 (Colombia)
#   argentina — Ley 25.326 (Argentina)
#   uk        — UK GDPR / DPA 2018 / DUAA 2025
#   ccpa      — CCPA/CPRA + regs. enero 2026 (California, EE.UU.)
#   todo      — Activa todas las jurisdicciones simultáneamente
#
# Uso — anonimización:
#   python anonimizar.py archivo.csv                  → genera archivo_anon.csv + archivo_anon.csv.key.json
#   python anonimizar.py *.docx --ley chile           → varios archivos
#   python anonimizar.py C:/exports/ --ley rgpd       → carpeta completa
#   python anonimizar.py datos.csv --salida limpio.csv → salida explícita
#   python anonimizar.py C:/exports/ --carpeta-salida C:/anon/
#   python anonimizar.py --lista-leyes
#
# Uso — restauración:
#   python anonimizar.py archivo_anon.csv --restaurar
#     → busca archivo_anon.csv.key.json en la misma carpeta → genera archivo_anon_restaurado.csv
#   python anonimizar.py archivo_anon.csv --restaurar --mapa otra_ruta.key.json
#     → usa el archivo de mapa indicado explícitamente
#
# Formatos soportados: .csv, .xlsx, .md, .docx
# Salida anonimización: [nombre]_anon.[ext] + [nombre]_anon.[ext].key.json
# Salida restauración:  [nombre]_restaurado.[ext]
# =============================================================================

import os
import re
import sys
import json
import argparse
import datetime

# =============================================================================
# PARTE 0: Argumentos de línea de comandos
# =============================================================================

JURISDICCIONES_DISPONIBLES = ["rgpd", "chile", "brasil", "mexico", "colombia", "argentina", "uk", "ccpa"]

parser = argparse.ArgumentParser(
    description="Anonimizador multi-jurisdiccional de datos personales.",
    formatter_class=argparse.RawTextHelpFormatter,
    epilog=(
        "Ejemplos:\n"
        "  python anonimizar.py informe.docx\n"
        "  python anonimizar.py datos.csv clientes.xlsx --ley chile\n"
        "  python anonimizar.py C:/exports/ --ley rgpd\n"
        "  python anonimizar.py datos.csv --salida datos_limpio.csv --ley rgpd\n"
        "  python anonimizar.py C:/exports/ --carpeta-salida C:/anon/ --ley todo\n"
        "  python anonimizar.py datos_anon.csv --restaurar\n"
        "  python anonimizar.py datos_anon.csv --restaurar --mapa datos_anon.csv.key.json\n"
    )
)
parser.add_argument(
    "entradas", nargs="*", metavar="ARCHIVO_O_CARPETA",
    help="Archivos o carpetas a procesar. Acepta varios a la vez."
)
parser.add_argument(
    "--ley", nargs="+", default=["rgpd"],
    metavar="LEY",
    help=(
        "Jurisdicción(es) a aplicar:\n"
        "  rgpd       UE — RGPD/GDPR 2016/679\n"
        "  chile      Chile — Ley 21.719 / 19.628\n"
        "  brasil     Brasil — LGPD Lei 13.709/2018\n"
        "  mexico     México — LFPDPPP 2010/2025\n"
        "  colombia   Colombia — Ley 1581/2012\n"
        "  argentina  Argentina — Ley 25.326\n"
        "  uk         UK GDPR / DPA 2018 / DUAA 2025\n"
        "  ccpa       California — CCPA/CPRA\n"
        "  todo       Activa todas las anteriores\n"
    )
)
parser.add_argument(
    "--salida", metavar="ARCHIVO",
    help="Ruta de salida explícita (solo válido con un único archivo de entrada)."
)
parser.add_argument(
    "--carpeta-salida", metavar="CARPETA",
    help="Carpeta de destino para todos los archivos procesados."
)
parser.add_argument(
    "--lista-leyes", action="store_true",
    help="Muestra las jurisdicciones disponibles y termina."
)
parser.add_argument(
    "--restaurar", action="store_true",
    help=(
        "Modo restauración: reemplaza los tokens numerados por los valores originales.\n"
        "Requiere el archivo .key.json generado durante la anonimización."
    )
)
parser.add_argument(
    "--mapa", metavar="ARCHIVO_KEY",
    help=(
        "Ruta explícita al archivo .key.json de mapa.\n"
        "Por defecto se busca [archivo_anon].[ext].key.json en la misma carpeta."
    )
)
args = parser.parse_args()

if args.lista_leyes:
    print("Jurisdicciones disponibles:")
    for ley in JURISDICCIONES_DISPONIBLES:
        print(f"  {ley}")
    sys.exit(0)

# =============================================================================
# PARTE 0b: Modo restauración — inicialización temprana (sin NLP)
# =============================================================================
# En modo restauración no se necesita spaCy ni Presidio.
# Se importan solo los módulos de E/S de archivo.

if args.restaurar:
    import pandas as pd
    from docx import Document

    def restaurar_texto(texto: str, mapa: dict) -> str:
        if not texto or not texto.strip():
            return texto
        # Ordenar por longitud descendente para evitar sustituciones parciales
        for token, original in sorted(mapa.items(), key=lambda x: len(x[0]), reverse=True):
            texto = texto.replace(token, original)
        return texto

    def cargar_mapa(ruta_mapa: str) -> dict:
        with open(ruta_mapa, "r", encoding="utf-8") as f:
            datos = json.load(f)
        return datos["mapa"]

    def _mapas_en_carpeta(carpeta: str) -> list:
        """Lista los archivos .key.json presentes en una carpeta."""
        if not os.path.isdir(carpeta):
            return []
        return sorted(
            os.path.join(carpeta, n)
            for n in os.listdir(carpeta)
            if n.lower().endswith(".key.json")
        )

    def ruta_mapa_para(ruta_entrada: str, mapa_explicito) -> str:
        """
        Resuelve qué .key.json usar para un archivo a restaurar.
        Orden de prioridad:
          1. --mapa explícito (se aplica a todas las entradas).
          2. [archivo].key.json exacto, junto al archivo (caso anonimización directa).
          3. Si en la carpeta hay un único .key.json, se usa ese (caso archivo
             devuelto por una IA con nombre distinto pero junto a su mapa original).
          4. Si hay varios .key.json y ninguno coincide por nombre → error pidiendo --mapa.
        """
        if mapa_explicito:
            return mapa_explicito
        base = os.path.abspath(ruta_entrada)
        # 2 — coincidencia exacta por nombre
        candidato = base + ".key.json"
        if os.path.isfile(candidato):
            return candidato
        # 3 / 4 — buscar en la carpeta del archivo
        carpeta = os.path.dirname(base)
        mapas = _mapas_en_carpeta(carpeta)
        if len(mapas) == 1:
            return mapas[0]
        if len(mapas) == 0:
            raise FileNotFoundError(
                f"No se encontró ningún .key.json para '{ruta_entrada}'.\n"
                f"Buscado: {candidato} y archivos .key.json en {carpeta}\n"
                f"Usa --mapa para indicar la ruta del mapa."
            )
        raise FileNotFoundError(
            f"Hay {len(mapas)} archivos .key.json en {carpeta} y ninguno coincide "
            f"con el nombre de '{os.path.basename(ruta_entrada)}'.\n"
            f"Indica cuál usar con --mapa para evitar mezclar tokens de archivos distintos."
        )

    def ruta_salida_restaurada_para(ruta_entrada: str) -> str:
        if args.salida:
            return args.salida
        nombre_base, ext = os.path.splitext(os.path.basename(ruta_entrada))
        nombre_restaurado = f"{nombre_base}_restaurado{ext}"
        if args.carpeta_salida:
            os.makedirs(args.carpeta_salida, exist_ok=True)
            return os.path.join(args.carpeta_salida, nombre_restaurado)
        return os.path.join(os.path.dirname(os.path.abspath(ruta_entrada)), nombre_restaurado)

    def restaurar_csv(ruta_entrada: str, ruta_salida: str, mapa: dict):
        df = pd.read_csv(ruta_entrada, dtype=str)
        df = df.map(lambda c: restaurar_texto(c, mapa) if isinstance(c, str) else c)
        df.to_csv(ruta_salida, index=False)
        print(f"  CSV restaurado: {ruta_salida}")

    def restaurar_xlsx(ruta_entrada: str, ruta_salida: str, mapa: dict):
        hojas = pd.read_excel(ruta_entrada, sheet_name=None, dtype=str)
        hojas_rest = {}
        for nombre, df in hojas.items():
            df = df.map(lambda c: restaurar_texto(c, mapa) if isinstance(c, str) else c)
            hojas_rest[nombre] = df
        with pd.ExcelWriter(ruta_salida, engine="openpyxl") as writer:
            for nombre, df in hojas_rest.items():
                df.to_excel(writer, sheet_name=nombre, index=False)
        print(f"  XLSX restaurado: {ruta_salida}")

    def restaurar_md(ruta_entrada: str, ruta_salida: str, mapa: dict):
        with open(ruta_entrada, "r", encoding="utf-8") as f:
            lineas = f.readlines()
        lineas_rest = [restaurar_texto(l, mapa) for l in lineas]
        with open(ruta_salida, "w", encoding="utf-8") as f:
            f.writelines(lineas_rest)
        print(f"  MD restaurado: {ruta_salida}")

    def restaurar_docx(ruta_entrada: str, ruta_salida: str, mapa: dict):
        doc = Document(ruta_entrada)
        for parrafo in doc.paragraphs:
            for run in parrafo.runs:
                run.text = restaurar_texto(run.text, mapa)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        for run in parrafo.runs:
                            run.text = restaurar_texto(run.text, mapa)
        doc.save(ruta_salida)
        print(f"  DOCX restaurado: {ruta_salida}")

    RESTAURADORES = {
        ".csv":  restaurar_csv,
        ".xlsx": restaurar_xlsx,
        ".md":   restaurar_md,
        ".docx": restaurar_docx,
    }

    def recopilar_archivos_restaurar(entradas: list) -> list:
        """
        Expande archivos y carpetas en rutas de archivos a restaurar.
        Excluye los propios .key.json y los archivos ya restaurados (_restaurado).
        En carpetas, solo recorre el primer nivel.
        """
        rutas = []
        for entrada in entradas:
            entrada = os.path.abspath(entrada)
            if os.path.isdir(entrada):
                for nombre in sorted(os.listdir(entrada)):
                    ruta = os.path.join(entrada, nombre)
                    if os.path.isfile(ruta):
                        rutas.append(ruta)
            elif os.path.isfile(entrada):
                rutas.append(entrada)
            else:
                print(f"  [AVISO] No encontrado: {entrada}")
        # Filtrar mapas y salidas previas
        filtradas = []
        for r in rutas:
            base = os.path.basename(r).lower()
            ext = os.path.splitext(r)[1].lower()
            if base.endswith(".key.json"):
                continue
            if os.path.splitext(os.path.basename(r))[0].endswith("_restaurado"):
                continue
            if ext not in RESTAURADORES:
                print(f"  [OMITIDO] {r} (extensión no soportada)")
                continue
            filtradas.append(r)
        return filtradas

    if not args.entradas:
        parser.print_help()
        sys.exit(0)

    archivos_rest = recopilar_archivos_restaurar(args.entradas)

    if args.salida and len(archivos_rest) > 1:
        print("[ERROR] --salida solo es válido con un único archivo de entrada.")
        sys.exit(1)

    if not archivos_rest:
        print("No se encontraron archivos para restaurar.")
        sys.exit(0)

    print(f"Modo: RESTAURACIÓN — archivos a procesar: {len(archivos_rest)}\n")
    for ruta_entrada in archivos_rest:
        extension = os.path.splitext(ruta_entrada)[1].lower()
        print(f"Restaurando: {ruta_entrada}")
        try:
            ruta_mapa = ruta_mapa_para(ruta_entrada, args.mapa)
            mapa = cargar_mapa(ruta_mapa)
            print(f"  Mapa cargado: {ruta_mapa} ({len(mapa)} tokens)")
            ruta_salida = ruta_salida_restaurada_para(ruta_entrada)
            RESTAURADORES[extension](ruta_entrada, ruta_salida, mapa)
        except Exception as e:
            print(f"  [ERROR] {e}")

    sys.exit(0)

# =============================================================================
# PARTE 1: Modelos de lenguaje — spaCy multilingüe
# =============================================================================

import spacy
import pandas as pd
from docx import Document
from langdetect import detect, LangDetectException

from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_analyzer.nlp_engine import NlpEngineProvider

leyes_activas = set(JURISDICCIONES_DISPONIBLES if "todo" in args.ley else args.ley)
leyes_invalidas = leyes_activas - set(JURISDICCIONES_DISPONIBLES)
if leyes_invalidas:
    print(f"[ERROR] Jurisdicción(es) desconocida(s): {', '.join(leyes_invalidas)}")
    print(f"Valores válidos: {', '.join(JURISDICCIONES_DISPONIBLES)}, todo")
    sys.exit(1)

print(f"Jurisdicciones activas: {', '.join(sorted(leyes_activas))}\n")

MODELOS_IDIOMA = {
    "es": "es_core_news_lg",
    "en": "en_core_web_lg",
    "fr": "fr_core_news_lg",
    "de": "de_core_news_lg",
    "it": "it_core_news_lg",
    "pt": "pt_core_news_lg",
    "nl": "nl_core_news_lg",
}

print("Comprobando modelos de idioma disponibles:")
modelos_cargados = []
idiomas_disponibles = []

for lang_code, model_name in MODELOS_IDIOMA.items():
    try:
        spacy.load(model_name)
        modelos_cargados.append({"lang_code": lang_code, "model_name": model_name})
        idiomas_disponibles.append(lang_code)
        print(f"  OK  {lang_code} ({model_name})")
    except OSError:
        print(f"  --  {lang_code} ({model_name}) no instalado — omitido")

if not modelos_cargados:
    raise RuntimeError(
        "No hay ningún modelo de spaCy instalado.\n"
        "Instala al menos uno con: python -m spacy download es_core_news_lg"
    )

idioma_por_defecto = idiomas_disponibles[0]

provider = NlpEngineProvider(nlp_configuration={
    "nlp_engine_name": "spacy",
    "models": modelos_cargados
})
nlp_engine = provider.create_engine()
analyzer = AnalyzerEngine(
    nlp_engine=nlp_engine,
    supported_languages=idiomas_disponibles
)

print(f"\nMotor listo. Idiomas activos: {', '.join(idiomas_disponibles)}\n")

# =============================================================================
# PARTE 2: Mapa de etiquetas y entidades universales
# =============================================================================

ETIQUETAS = {
    "PERSON":          "PERSONA",
    "EMAIL_ADDRESS":   "EMAIL",
    "PHONE_NUMBER":    "TELEFONO",
    "IBAN_CODE":       "IBAN",
    "CREDIT_CARD":     "TARJETA",
    "DATE_TIME":       "FECHA",
    "IP_ADDRESS":      "IP",
    "LOCATION":        "UBICACION",
    "DNI_NIE":         "DNI-ES",
    "RUT_CL":          "RUT-CL",
    "CPF_BR":          "CPF-BR",
    "CNPJ_BR":         "CNPJ-BR",
    "CURP_MX":         "CURP-MX",
    "RFC_MX":          "RFC-MX",
    "NIT_CO":          "NIT-CO",
    "CUIT_AR":         "CUIT-AR",
    "DNI_AR":          "DNI-AR",
    "NINO_UK":         "NINO-UK",
    "SSN_US":          "SSN-US",
    "DL_US":           "DL-US",
}

ENTIDADES_BASE = [
    "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER",
    "IBAN_CODE", "CREDIT_CARD", "DATE_TIME", "IP_ADDRESS", "LOCATION",
]

FALSOS_POSITIVOS_PERSONA = {
    "hola", "buenos", "buenas", "días", "tardes", "noches",
    "estimado", "estimada", "estimados", "estimadas",
    "querido", "querida", "saludos", "atentamente", "cordialmente",
    "señor", "señora", "señorita", "don", "doña",
    "presidente", "director", "gerente", "cliente", "usuario",
}

# =============================================================================
# PARTE 3: Reconocedores por jurisdicción
# =============================================================================

def _registrar_para_todos(name_prefix, entity, patterns):
    for lang in idiomas_disponibles:
        name = f"{name_prefix}_{lang}" if name_prefix else f"{entity}_{lang}"
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=name,
            patterns=patterns,
            supported_language=lang,
            supported_entity=entity,
        ))


def activar_rgpd():
    analyzer.registry.add_recognizer(PatternRecognizer(
        name="DNI_NIE_ES",
        patterns=[
            Pattern("DNI", r"\b\d{8}[A-HJ-NP-TV-Za-hj-np-tv-z]\b", 0.85),
            Pattern("NIE", r"\b[XYZxyz]\d{7}[A-HJ-NP-TV-Za-hj-np-tv-z]\b", 0.85),
        ],
        supported_language="es",
        supported_entity="DNI_NIE",
    ))
    patrones_tel_es = [
        Pattern("Tel_ES_intl",  r"(?<!\d)(?:\+34|0034)[\s.\-]?[6-9]\d{2}[\s.\-]?\d{3}[\s.\-]?\d{3}\b", 0.95),
        Pattern("Movil_ES",     r"\b[67]\d{2}[\s.\-]?\d{3}[\s.\-]?\d{3}\b", 0.80),
        Pattern("Fijo_ES",      r"\b9\d{2}[\s.\-]?\d{3}[\s.\-]?\d{3}\b", 0.75),
        Pattern("Tel_intl_gen", r"(?<!\d)\+(?!34|0034)\d{1,3}[\s.\-]?\(?\d{1,4}\)?[\s.\-]?\d{2,4}[\s.\-]?\d{2,4}[\s.\-]?\d{0,4}\b", 0.85),
    ]
    _registrar_para_todos(None, "PHONE_NUMBER", patrones_tel_es)
    print("  [RGPD] Reconocedores DNI/NIE + teléfonos ES activados.")


def activar_chile():
    patrones_rut = [
        Pattern("RUT_puntos",  r"\b\d{1,2}\.\d{3}\.\d{3}-[\dKk]\b", 0.95),
        Pattern("RUT_sin_pts", r"\b\d{7,8}-[\dKk]\b",               0.85),
    ]
    _registrar_para_todos(None, "RUT_CL", patrones_rut)
    patrones_tel_cl = [
        Pattern("Tel_CL_movil", r"(?:\+56|0056)?[\s\-]?9[\s\-]?\d{4}[\s\-]?\d{4}\b", 0.90),
        Pattern("Tel_CL_fijo",  r"(?:\+56|0056)[\s\-]?[2-9]\d{7}\b", 0.85),
    ]
    _registrar_para_todos(None, "PHONE_NUMBER", patrones_tel_cl)
    print("  [Chile — Ley 21.719] Reconocedores RUT/RUN + teléfonos CL activados.")


def activar_brasil():
    patrones_cpf = [
        Pattern("CPF_puntos", r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b", 0.95),
        Pattern("CPF_sin",    r"\b\d{11}\b",                     0.60),
    ]
    patrones_cnpj = [
        Pattern("CNPJ_fmt", r"\b\d{2}\.\d{3}\.\d{3}\/\d{4}-\d{2}\b", 0.95),
        Pattern("CNPJ_sin", r"\b\d{14}\b",                            0.55),
    ]
    for lang in idiomas_disponibles:
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"CPF_BR_{lang}", patterns=patrones_cpf,
            supported_language=lang, supported_entity="CPF_BR",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"CNPJ_BR_{lang}", patterns=patrones_cnpj,
            supported_language=lang, supported_entity="CNPJ_BR",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"Telefono_BR_{lang}",
            patterns=[
                Pattern("Tel_BR_movil", r"(?:\+55|0055)?[\s\-]?\(?\d{2}\)?[\s\-]?9\d{4}[\s\-]?\d{4}\b", 0.90),
                Pattern("Tel_BR_fijo",  r"(?:\+55|0055)?[\s\-]?\(?\d{2}\)?[\s\-]?\d{4}[\s\-]?\d{4}\b", 0.80),
            ],
            supported_language=lang, supported_entity="PHONE_NUMBER",
        ))
    print("  [Brasil — LGPD] Reconocedores CPF, CNPJ + teléfonos BR activados.")


def activar_mexico():
    patrones_curp = [
        Pattern("CURP_MX",
            r"\b[A-Z]{4}\d{6}[HM][A-Z]{5}[B-DF-HJ-NP-TV-Z\d]\d\b",
            0.92),
    ]
    patrones_rfc = [
        Pattern("RFC_fisica", r"\b[A-Z&Ñ]{4}\d{6}[A-Z0-9]{3}\b", 0.85),
        Pattern("RFC_moral",  r"\b[A-Z&Ñ]{3}\d{6}[A-Z0-9]{3}\b", 0.80),
    ]
    for lang in idiomas_disponibles:
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"CURP_MX_{lang}", patterns=patrones_curp,
            supported_language=lang, supported_entity="CURP_MX",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"RFC_MX_{lang}", patterns=patrones_rfc,
            supported_language=lang, supported_entity="RFC_MX",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"Telefono_MX_{lang}",
            patterns=[
                Pattern("Tel_MX",
                    r"(?:\+52|0052)?[\s\-]?\(?\d{2,3}\)?[\s\-]?\d{3,4}[\s\-]?\d{4}\b",
                    0.80),
            ],
            supported_language=lang, supported_entity="PHONE_NUMBER",
        ))
    print("  [México — LFPDPPP] Reconocedores CURP, RFC + teléfonos MX activados.")


def activar_colombia():
    patrones_nit = [
        Pattern("NIT_CO", r"\b\d{9}-\d\b", 0.90),
    ]
    patrones_cc = [
        Pattern("CC_CO_contexto",
            r"(?:C\.?C\.?|cédula|documento)\s*:?\s*(\d{6,10})\b",
            0.85),
    ]
    for lang in idiomas_disponibles:
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"NIT_CO_{lang}", patterns=patrones_nit,
            supported_language=lang, supported_entity="NIT_CO",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"CC_CO_{lang}", patterns=patrones_cc,
            supported_language=lang, supported_entity="DNI_NIE",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"Telefono_CO_{lang}",
            patterns=[
                Pattern("Tel_CO_movil", r"(?:\+57|0057)?[\s\-]?3\d{2}[\s\-]?\d{3}[\s\-]?\d{4}\b", 0.90),
                Pattern("Tel_CO_fijo",  r"(?:\+57|0057)?[\s\-]?\(?\d{1,3}\)?[\s\-]?\d{3}[\s\-]?\d{4}\b", 0.75),
            ],
            supported_language=lang, supported_entity="PHONE_NUMBER",
        ))
    print("  [Colombia — Ley 1581] Reconocedores NIT, CC + teléfonos CO activados.")


def activar_argentina():
    patrones_dni_ar = [
        Pattern("DNI_AR_puntos", r"\b\d{2}\.\d{3}\.\d{3}\b", 0.85),
        Pattern("DNI_AR_sin",    r"\b[1-9]\d{6,7}\b",         0.55),
    ]
    patrones_cuit = [
        Pattern("CUIT_AR", r"\b\d{2}-\d{8}-\d\b", 0.95),
    ]
    for lang in idiomas_disponibles:
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"DNI_AR_{lang}", patterns=patrones_dni_ar,
            supported_language=lang, supported_entity="DNI_AR",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"CUIT_AR_{lang}", patterns=patrones_cuit,
            supported_language=lang, supported_entity="CUIT_AR",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"Telefono_AR_{lang}",
            patterns=[
                Pattern("Tel_AR",
                    r"(?:\+54|0054)?[\s\-]?\(?\d{2,4}\)?[\s\-]?1{0,1}\d{4}[\s\-]?\d{4}\b",
                    0.80),
            ],
            supported_language=lang, supported_entity="PHONE_NUMBER",
        ))
    print("  [Argentina — Ley 25.326] Reconocedores DNI, CUIT/CUIL + teléfonos AR activados.")


def activar_uk():
    patrones_nino = [
        Pattern("NINO_UK",
            r"\b(?!BG|GB|NK|KN|TN|NT|ZZ)[A-CEGHJ-PR-TW-Z]{2}[\s]?\d{2}[\s]?\d{2}[\s]?\d{2}[\s]?[A-D]\b",
            0.92),
    ]
    patrones_tel_uk = [
        Pattern("Tel_UK",
            r"(?:\+44|0044|0)[\s\-]?\(?\d{2,4}\)?[\s\-]?\d{3,4}[\s\-]?\d{3,4}\b",
            0.85),
    ]
    for lang in idiomas_disponibles:
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"NINO_UK_{lang}", patterns=patrones_nino,
            supported_language=lang, supported_entity="NINO_UK",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"Telefono_UK_{lang}", patterns=patrones_tel_uk,
            supported_language=lang, supported_entity="PHONE_NUMBER",
        ))
    print("  [UK — UK GDPR/DPA 2018] Reconocedores NINO + teléfonos UK activados.")


def activar_ccpa():
    patrones_ssn = [
        Pattern("SSN_US",   r"\b(?!000|666|9\d{2})\d{3}-(?!00)\d{2}-(?!0000)\d{4}\b", 0.95),
    ]
    patrones_dl = [
        Pattern("DL_CA_US", r"\b[A-Z]\d{7}\b", 0.70),
    ]
    patrones_tel_us = [
        Pattern("Tel_US",
            r"(?:\+1|001)?[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}\b",
            0.80),
    ]
    for lang in idiomas_disponibles:
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"SSN_US_{lang}", patterns=patrones_ssn,
            supported_language=lang, supported_entity="SSN_US",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"DL_US_{lang}", patterns=patrones_dl,
            supported_language=lang, supported_entity="DL_US",
        ))
        analyzer.registry.add_recognizer(PatternRecognizer(
            name=f"Telefono_US_{lang}", patterns=patrones_tel_us,
            supported_language=lang, supported_entity="PHONE_NUMBER",
        ))
    print("  [CCPA/CPRA — California] Reconocedores SSN, DL + teléfonos US activados.")


# =============================================================================
# PARTE 4: Activar reconocedores según jurisdicciones seleccionadas
# =============================================================================

ACTIVADORES = {
    "rgpd":      activar_rgpd,
    "chile":     activar_chile,
    "brasil":    activar_brasil,
    "mexico":    activar_mexico,
    "colombia":  activar_colombia,
    "argentina": activar_argentina,
    "uk":        activar_uk,
    "ccpa":      activar_ccpa,
}

print("Activando reconocedores:")
for ley in sorted(leyes_activas):
    ACTIVADORES[ley]()

ENTIDADES = ENTIDADES_BASE + [
    e for e in ETIQUETAS
    if e not in ENTIDADES_BASE and any(
        rec.supported_entities and e in rec.supported_entities
        for rec in analyzer.registry.recognizers
    )
]

# =============================================================================
# PARTE 5: Detección de exports de Screaming Frog
# =============================================================================

COLUMNAS_TEXTO_SF = {
    "title 1", "title 2",
    "meta description 1", "meta description 2",
    "h1-1", "h1-2", "h2-1", "h2-2", "h3-1", "h3-2",
    "meta keywords 1", "snippet",
}

def es_archivo_screaming_frog(df) -> bool:
    cabeceras = {c.lower() for c in df.columns}
    return {"address", "content type", "status code"}.issubset(cabeceras)

# =============================================================================
# PARTE 6: Mapa de anonimización — estado por archivo
# =============================================================================
# Tokens numerados por tipo: <PERSONA-1>, <PERSONA-2>, <EMAIL-1>, <DNI-ES-1>, etc.
# El mismo valor original siempre recibe el mismo token dentro del archivo.
# El mapa se guarda en [archivo_anon].key.json como JSON legible.

_mapa_token_a_original: dict = {}
_mapa_original_a_token: dict = {}
_contadores_tipo: dict = {}


def _reset_mapa():
    _mapa_token_a_original.clear()
    _mapa_original_a_token.clear()
    _contadores_tipo.clear()


def _token_para(entidad_tipo: str, valor_original: str) -> str:
    if valor_original in _mapa_original_a_token:
        return _mapa_original_a_token[valor_original]
    prefijo = ETIQUETAS.get(entidad_tipo, entidad_tipo)
    n = _contadores_tipo.get(prefijo, 0) + 1
    _contadores_tipo[prefijo] = n
    token = f"<{prefijo}-{n}>"
    _mapa_original_a_token[valor_original] = token
    _mapa_token_a_original[token] = valor_original
    return token


def _guardar_mapa(ruta_anon: str, archivo_origen: str):
    ruta_mapa = ruta_anon + ".key.json"
    datos = {
        "version": "2.2",
        "ley": sorted(leyes_activas),
        "fecha": datetime.datetime.now(datetime.timezone.utc).isoformat(),
        "archivo_origen": os.path.basename(archivo_origen),
        "advertencia": "Este archivo contiene datos personales originales. Trátalo con el mismo nivel de protección que el archivo fuente.",
        "mapa": _mapa_token_a_original.copy(),
    }
    with open(ruta_mapa, "w", encoding="utf-8") as f:
        json.dump(datos, f, ensure_ascii=False, indent=2)
    print(f"  Mapa guardado: {ruta_mapa} ({len(_mapa_token_a_original)} tokens)")

# =============================================================================
# PARTE 7: Función principal de anonimización de texto
# =============================================================================

def _detectar_idioma(texto: str) -> str:
    if len(texto.split()) < 4:
        return idioma_por_defecto
    try:
        idioma = detect(texto)
        return idioma if idioma in idiomas_disponibles else idioma_por_defecto
    except LangDetectException:
        return idioma_por_defecto


def _normalizar_mayusculas(texto: str) -> str:
    palabras = texto.split()
    if not palabras:
        return texto
    prop = sum(1 for p in palabras if p.isalpha() and p.isupper()) / len(palabras)
    return texto.title() if prop > 0.5 else texto


def anonimizar_texto(texto: str) -> str:
    if not texto or not texto.strip():
        return texto

    texto_para_analisis = _normalizar_mayusculas(texto)
    idioma = _detectar_idioma(texto_para_analisis)

    resultados = analyzer.analyze(
        text=texto_para_analisis,
        entities=ENTIDADES,
        language=idioma
    )

    resultados = [
        r for r in resultados
        if not (
            r.entity_type == "PERSON" and
            texto[r.start:r.end].lower() in FALSOS_POSITIVOS_PERSONA
        )
    ]

    personas = [r for r in resultados if r.entity_type == "PERSON"]
    resultados = [
        r for r in resultados
        if r.entity_type != "LOCATION" or any(
            abs(r.start - p.end) <= 120 or abs(p.start - r.end) <= 120
            for p in personas
        )
    ]

    for r in resultados:
        if r.entity_type == "EMAIL_ADDRESS":
            r.score = 1.0

    resultados_sin_solapamiento = []
    for r in sorted(resultados, key=lambda x: x.score, reverse=True):
        solapado = any(
            r.start < e.end and r.end > e.start
            for e in resultados_sin_solapamiento
        )
        if not solapado:
            resultados_sin_solapamiento.append(r)

    resultados = sorted(resultados_sin_solapamiento, key=lambda r: r.start, reverse=True)

    texto_anon = texto
    for resultado in resultados:
        valor_original = texto[resultado.start:resultado.end]
        token = _token_para(resultado.entity_type, valor_original)
        texto_anon = texto_anon[:resultado.start] + token + texto_anon[resultado.end:]

    return texto_anon

# =============================================================================
# PARTE 8: Procesadores por formato de archivo
# =============================================================================

def procesar_csv(ruta_entrada: str, ruta_salida: str):
    df = pd.read_csv(ruta_entrada, dtype=str)
    if es_archivo_screaming_frog(df):
        print("  Detectado formato Screaming Frog — procesando solo columnas de texto libre.")
        for columna in df.columns:
            if columna.lower() in COLUMNAS_TEXTO_SF:
                df[columna] = df[columna].apply(
                    lambda c: anonimizar_texto(c) if isinstance(c, str) else c
                )
    else:
        df = df.map(lambda c: anonimizar_texto(c) if isinstance(c, str) else c)
    df.to_csv(ruta_salida, index=False)
    print(f"  CSV guardado: {ruta_salida}")


def procesar_xlsx(ruta_entrada: str, ruta_salida: str):
    hojas = pd.read_excel(ruta_entrada, sheet_name=None, dtype=str)
    hojas_anon = {}
    for nombre, df in hojas.items():
        df = df.map(lambda c: anonimizar_texto(c) if isinstance(c, str) else c)
        hojas_anon[nombre] = df
    with pd.ExcelWriter(ruta_salida, engine="openpyxl") as writer:
        for nombre, df in hojas_anon.items():
            df.to_excel(writer, sheet_name=nombre, index=False)
    print(f"  XLSX guardado: {ruta_salida}")


def procesar_md(ruta_entrada: str, ruta_salida: str):
    with open(ruta_entrada, "r", encoding="utf-8") as f:
        lineas = f.readlines()
    lineas_anon = [anonimizar_texto(l) for l in lineas]
    with open(ruta_salida, "w", encoding="utf-8") as f:
        f.writelines(lineas_anon)
    print(f"  MD guardado: {ruta_salida}")


def procesar_docx(ruta_entrada: str, ruta_salida: str):
    doc = Document(ruta_entrada)
    for parrafo in doc.paragraphs:
        for run in parrafo.runs:
            run.text = anonimizar_texto(run.text)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        run.text = anonimizar_texto(run.text)
    doc.save(ruta_salida)
    print(f"  DOCX guardado: {ruta_salida}")

# =============================================================================
# PARTE 9: Resolución de rutas y bucle principal
# =============================================================================

PROCESADORES = {
    ".csv":  procesar_csv,
    ".xlsx": procesar_xlsx,
    ".md":   procesar_md,
    ".docx": procesar_docx,
}


def ruta_salida_para(ruta_entrada: str) -> str:
    if args.salida:
        return args.salida
    nombre, ext = os.path.splitext(os.path.basename(ruta_entrada))
    nombre_anon = f"{nombre}_anon{ext}"
    if args.carpeta_salida:
        os.makedirs(args.carpeta_salida, exist_ok=True)
        return os.path.join(args.carpeta_salida, nombre_anon)
    return os.path.join(os.path.dirname(os.path.abspath(ruta_entrada)), nombre_anon)


def recopilar_archivos(entradas: list) -> list:
    rutas = []
    for entrada in entradas:
        entrada = os.path.abspath(entrada)
        if os.path.isdir(entrada):
            for nombre in os.listdir(entrada):
                ruta = os.path.join(entrada, nombre)
                if os.path.isfile(ruta):
                    rutas.append(ruta)
        elif os.path.isfile(entrada):
            rutas.append(entrada)
        else:
            print(f"  [AVISO] No encontrado: {entrada}")
    return rutas


if not args.entradas:
    parser.print_help()
    sys.exit(0)

if args.salida and len(args.entradas) > 1:
    print("[ERROR] --salida solo es válido con un único archivo de entrada.")
    sys.exit(1)

archivos = recopilar_archivos(args.entradas)

if not archivos:
    print("No se encontraron archivos para procesar.")
    sys.exit(0)

print(f"Archivos a procesar: {len(archivos)}\n")
for ruta_entrada in archivos:
    _, extension = os.path.splitext(ruta_entrada)
    extension = extension.lower()
    if extension not in PROCESADORES:
        print(f"  [OMITIDO] {ruta_entrada} (extensión no soportada)")
        continue
    ruta_salida = ruta_salida_para(ruta_entrada)
    print(f"Procesando: {ruta_entrada}")
    _reset_mapa()
    try:
        PROCESADORES[extension](ruta_entrada, ruta_salida)
        _guardar_mapa(ruta_salida, ruta_entrada)
    except Exception as e:
        print(f"  [ERROR] {e}")
